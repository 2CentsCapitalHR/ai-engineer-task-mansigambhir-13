#!/usr/bin/env python3
"""
ADGM Corporate Agent - Test Sample Script
Demonstrates the core functionality with sample documents
"""

import json
import os
from datetime import datetime
from pathlib import Path

# Import main classes (assuming they're in separate modules)
try:
    from main import ADGMCorporateAgent, ADGMKnowledgeBase, DocumentProcessor
except ImportError:
    print("Please ensure main.py is in the same directory")
    exit(1)

def create_sample_documents():
    """Create sample .docx documents for testing"""
    try:
        from docx import Document
        
        # Sample Articles of Association with issues
        doc1 = Document()
        doc1.add_heading('Articles of Association', 0)
        doc1.add_paragraph("""
        ARTICLES OF ASSOCIATION
        
        1. Company Name: [Company Name TBD]
        
        2. Jurisdiction: These articles shall be governed by UAE Federal Law.
        
        3. Share Capital: The authorized share capital is XXX AED.
        
        4. Directors: The company shall have [number] directors.
        
        5. Registered Office: [Address to be determined]
        
        Signature: ________________
        Director: 
        """)
        
        doc1.save("sample_aoa_with_issues.docx")
        
        # Sample Employment Contract with fewer issues
        doc2 = Document()
        doc2.add_heading('Employment Contract', 0)
        doc2.add_paragraph("""
        EMPLOYMENT CONTRACT
        
        This employment contract is governed by ADGM Employment Regulations.
        
        Employee: John Doe
        Position: Senior Manager
        Salary: AED 15,000 per month
        
        Jurisdiction: This contract shall be subject to ADGM Courts.
        
        Signed: ________________
        Employee: John Doe
        
        Signed: ________________
        Employer: ABC Company Limited
        """)
        
        doc2.save("sample_employment_contract.docx")
        
        print("✅ Sample documents created:")
        print("  - sample_aoa_with_issues.docx")
        print("  - sample_employment_contract.docx")
        
        return ["sample_aoa_with_issues.docx", "sample_employment_contract.docx"]
        
    except ImportError:
        print("❌ python-docx not available. Cannot create sample documents.")
        return []

def test_document_analysis():
    """Test the document analysis functionality"""
    print("\n🔍 Testing ADGM Corporate Agent...")
    
    # Initialize the agent
    agent = ADGMCorporateAgent()
    
    # Create sample documents
    sample_files = create_sample_documents()
    
    if not sample_files:
        print("❌ No sample documents available for testing")
        return
    
    # Test knowledge base
    print("\n📚 Testing Knowledge Base...")
    kb = agent.knowledge_base
    
    # Test requirements lookup
    incorporation_reqs = kb.get_requirements_for_process("company_incorporation")
    print(f"📋 Incorporation Requirements ({len(incorporation_reqs)} items):")
    for req in incorporation_reqs:
        print(f"  • {req}")
    
    # Test jurisdiction checking
    test_text = "This contract is governed by UAE Federal Law and Dubai Courts."
    jurisdiction_issues = kb.check_jurisdiction_compliance(test_text)
    print(f"\n⚖️ Jurisdiction Issues Found: {len(jurisdiction_issues)}")
    for issue in jurisdiction_issues:
        print(f"  • {issue}")
    
    # Test document processing
    print("\n📄 Testing Document Processing...")
    processor = agent.processor
    
    for file_path in sample_files:
        if os.path.exists(file_path):
            print(f"\n--- Processing: {file_path} ---")
            
            # Extract text
            text = processor.extract_text_from_docx(file_path)
            print(f"📝 Text Length: {len(text)} characters")
            
            # Identify document type
            doc_type = processor.identify_document_type(text, file_path)
            print(f"📋 Document Type: {doc_type}")
            
            # Detect red flags
            issues = processor.detect_red_flags(text, doc_type)
            print(f"🚩 Red Flags Found: {len(issues)}")
            
            for i, issue in enumerate(issues[:3], 1):  # Show first 3
                print(f"  {i}. {issue.severity}: {issue.issue}")
                print(f"     💡 Suggestion: {issue.suggestion}")
                if issue.adgm_reference:
                    print(f"     📚 Reference: {issue.adgm_reference}")

def test_full_analysis():
    """Test the complete analysis workflow"""
    print("\n🎯 Testing Complete Analysis Workflow...")
    
    agent = ADGMCorporateAgent()
    sample_files = []
    
    # Check for existing sample files
    for filename in ["sample_aoa_with_issues.docx", "sample_employment_contract.docx"]:
        if os.path.exists(filename):
            # Create mock file object
            class MockFile:
                def __init__(self, name):
                    self.name = name
            sample_files.append(MockFile(filename))
    
    if not sample_files:
        print("❌ No sample files found. Run create_sample_documents() first.")
        return
    
    # Run full analysis
    result, status_msg = agent.analyze_documents(sample_files)
    
    if result:
        print("✅ Analysis completed successfully!")
        print(f"\n📊 Results Summary:")
        print(f"  Process: {result.process}")
        print(f"  Documents: {result.documents_uploaded}/{result.required_documents}")
        print(f"  Missing: {len(result.missing_documents)} documents")
        print(f"  Issues: {len(result.issues_found)} found")
        print(f"  Compliance Score: {result.compliance_score:.1f}/100")
        
        print(f"\n📋 Status Message:")
        print(status_msg)
        
        # Save results to JSON
        output_file = f"test_results_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
        with open(output_file, 'w') as f:
            json.dump({
                'result': {
                    'process': result.process,
                    'documents_uploaded': result.documents_uploaded,
                    'required_documents': result.required_documents,
                    'missing_documents': result.missing_documents,
                    'issues_found': [
                        {
                            'document': issue.document,
                            'section': issue.section,
                            'issue': issue.issue,
                            'severity': issue.severity,
                            'suggestion': issue.suggestion,
                            'adgm_reference': issue.adgm_reference
                        } for issue in result.issues_found
                    ],
                    'compliance_score': result.compliance_score
                },
                'status_message': status_msg,
                'timestamp': datetime.now().isoformat()
            }, indent=2)
        
        print(f"\n💾 Results saved to: {output_file}")
    else:
        print(f"❌ Analysis failed: {status_msg}")

def test_checklist_verification():
    """Test the document checklist verification feature"""
    print("\n📝 Testing Document Checklist Verification...")
    
    kb = ADGMKnowledgeBase()
    
    # Test different process types
    processes = ["company_incorporation", "licensing_requirements", "employment_requirements"]
    
    for process in processes:
        requirements = kb.get_requirements_for_process(process)
        print(f"\n📋 {process.replace('_', ' ').title()}:")
        print(f"   Required Documents: {len(requirements)}")
        for req in requirements:
            print(f"   • {req}")
    
    # Test missing document detection
    print("\n🔍 Testing Missing Document Detection...")
    
    # Simulate uploaded documents
    uploaded_docs = [
        "Articles of Association",
        "Memorandum of Association", 
        "Board Resolution"
    ]
    
    required_docs = kb.get_requirements_for_process("company_incorporation")
    
    missing = []
    for required in required_docs:
        found = any(req_word in uploaded.lower() 
                   for uploaded in uploaded_docs 
                   for req_word in required.lower().split())
        if not found:
            missing.append(required)
    
    print(f"📤 Uploaded: {uploaded_docs}")
    print(f"❌ Missing: {missing}")
    print(f"📊 Completion: {len(uploaded_docs)}/{len(required_docs)} documents")

def run_comprehensive_test():
    """Run all tests in sequence"""
    print("🚀 ADGM Corporate Agent - Comprehensive Test Suite")
    print("=" * 60)
    
    try:
        # Test 1: Document analysis
        test_document_analysis()
        
        # Test 2: Checklist verification  
        test_checklist_verification()
        
        # Test 3: Full analysis workflow
        test_full_analysis()
        
        print("\n" + "=" * 60)
        print("✅ All tests completed successfully!")
        print("\n🎯 Next Steps:")
        print("  1. Run the main application: python main.py")
        print("  2. Upload the generated sample documents")
        print("  3. Review the analysis results and compliance scores")
        print("  4. Check the reviewed documents with inline comments")
        
    except Exception as e:
        print(f"\n❌ Test failed with error: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    run_comprehensive_test()